from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Aashi_Dreams_Project_Origin_and_Prompt_Timeline.docx"


INK = RGBColor(24, 30, 45)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(84, 92, 110)
FILL = "E8EEF5"
LIGHT_FILL = "F7F9FC"
BORDER = "AEB9C8"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = color or INK
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
      borders = OxmlElement("w:tblBorders")
      tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def set_table_widths(table, widths: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(subtitle)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(5)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = BLUE if level <= 2 else DARK_BLUE


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.style = doc.styles["Normal"]
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = INK


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    set_table_borders(table)
    set_table_widths(table, [1.8, 4.7])
    for idx, (key, value) in enumerate(rows):
        set_cell_text(table.cell(idx, 0), key, bold=True, color=DARK_BLUE)
        set_cell_text(table.cell(idx, 1), value)
        set_cell_shading(table.cell(idx, 0), FILL)
        if idx % 2 == 1:
            set_cell_shading(table.cell(idx, 1), LIGHT_FILL)
    doc.add_paragraph()


def add_timeline_table(doc: Document, rows: list[tuple[str, str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    set_table_borders(table)
    set_table_widths(table, [0.65, 2.05, 2.0, 1.8])
    headers = ["Stage", "Prompt / Ask", "Motive Behind It", "Product Direction"]
    for idx, header in enumerate(headers):
        set_cell_text(table.cell(0, idx), header, bold=True, color=DARK_BLUE)
        set_cell_shading(table.cell(0, idx), FILL)
    for stage, prompt, motive, product in rows:
        cells = table.add_row().cells
        values = [stage, prompt, motive, product]
        for idx, value in enumerate(values):
            set_cell_text(cells[idx], value)
            if int(stage.split(".")[0]) % 2 == 0:
                set_cell_shading(cells[idx], LIGHT_FILL)
    doc.add_paragraph()


def build_doc() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    add_title(
        doc,
        "Aashi Dreams: Project Origin and Prompt Timeline",
        "How the idea moved from a personal AI twin into a product-ready admissions intelligence engine.",
    )

    add_key_value_table(
        doc,
        [
            ("Created for", "Anushka Navin Kumar"),
            ("Current product name", "Aashi Dreams Intelligence Engine"),
            ("Original core idea", "A living AI version of the student that evolves from school to university admission, career planning, and placement."),
            ("Main motive", "Build a personal competitive advantage first, then package it into a premium B2B tool for coaching centers."),
            ("Current route", "http://127.0.0.1:5173/ and partner demos through ?partner=karnal or ?partner=zenith"),
        ],
    )

    add_heading(doc, "1. Core Motive", 1)
    add_body(
        doc,
        "The project did not start as a normal study-abroad tracker. The motive was to create a future operating system: a tool that continuously understands Anushka's profile, predicts her best education and career paths, and tells her what to improve next.",
    )
    for item in [
        "Make the system personal first: Anushka's real profile, research, startup, IELTS, scholarships, universities, and career direction.",
        "Use the research paper and Mentally Prepare as hidden advantages, not side activities.",
        "Turn uncertain admission advice into structured probability, roadmap, evidence, and action.",
        "Go beyond admissions into ROI, visa risk, alumni outcomes, career paths, and scholarship survival.",
        "Eventually make the product white-label and demo-ready for coaching institutes.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. Prompt Timeline", 1)
    add_body(doc, "This table lists the major prompts and the motive behind each stage. The wording is cleaned slightly for readability while preserving the original intent.")
    timeline = [
        ("1. Idea spark", "Do not build another application tracker. Build the Student Digital Twin.", "Shift from a checklist app to a living AI profile that evolves with the student.", "Career Simulation Engine, AI Career DNA, Life Path Simulator."),
        ("2. Naming", "It is not related to Stride. Its name will be Aashi Dreams.", "Separate the product identity from other brands and make it personally owned.", "Aashi Dreams became the product identity."),
        ("3. First build", "Start building Aashi Dreams, a living AI version. Right now I am making it for myself only.", "Begin with a single-user personal version, not a large public platform.", "Personal dashboard, profile, university cards, roadmap logic."),
        ("4. Trust gap", "College list is limited. I want college rank and acceptance rate. I do not like the color; make it elegant.", "Make the app feel credible and premium, with better data and design authority.", "Rank fields, acceptance-rate fields, elegant UI pass."),
        ("5. Geography expansion", "Next I also want USA colleges.", "Do not restrict the system to UK/Europe; include high-value USA options.", "USA university seed data and USA-specific route logic."),
        ("6. Asia and live updates", "Asia countries also. I want it updated whenever news, forms, ROI, or WHO-type changes match.", "Add wider geography and a future radar for changing opportunities.", "Asia countries, Opportunity Radar, official sources, update-oriented thinking."),
        ("7. Visual direction", "Use this kind of design with good and nice font.", "Make the product visually memorable instead of looking like a plain school portal.", "Warmer Aashi Dreams visual language, better typography, polished cards."),
        ("8. College detail", "When I click the college, I want a full roadmap and about the college also.", "Every university must become an actionable plan, not just a list item.", "University detail modal, roadmap phases, requirements, visa/scholarship notes."),
        ("9. Bug pressure", "Not working. Solve bug. Use your power to give right information. College rank is not right.", "Reliability and data accuracy became more important than feature quantity.", "Bug fixes, source notes, rank confidence wording."),
        ("10. Full-stack architecture", "Build Anushka OS in Replit: React, Express, SQLite, Drizzle, CRUD modules, AI endpoint, seed data.", "Turn the idea into a real application with database-backed sections.", "Full-stack app architecture, API routes, seed data, dashboard KPIs."),
        ("11. Missing strategic modules", "Add Professor CRM, Profile Delta, Network Intelligence, and Research Reputation Tracking.", "Elite admissions are won through faculty fit, gaps, network, and research credibility.", "Faculty Intelligence, Profile Delta, People Tracker, Research Reputation."),
        ("12. Running and fixing", "Run it. It has so many errors; see and tell. Next.", "Move from concept to working local product and debug continuously.", "Release checks, runtime fixes, route cleanup."),
        ("13. Real applicant evidence", "This I got from one person who got admission.", "Use real applicant advice: LORs, scholarships, CGPA, startup positioning, self-directed applications.", "LOR importance, scholarship competitiveness, startup/SOP alignment."),
        ("14. Scholarship seriousness", "You need to research scholarship tracking as it is intense.", "Scholarships are not optional; funding timing can decide whether admission is usable.", "Scholarship Tracker, official links, probability and action notes."),
        ("15. Official research depth", "Use UCAS and LSE links as reference and find more. Go deeper into college websites, scholarships, internships.", "Replace generic internet advice with official-source intelligence.", "Official Sources page, research-backed university and scholarship detail."),
        ("16. Risk diversification", "USA, Asia, and more Europe. Europe looks risky; focus on USA also.", "Build options across regions so one geopolitical or visa risk does not break the plan.", "USA/Asia/Europe expansion, visa and ROI comparison."),
        ("17. Requirement quality", "The NYU example explains requirements nicely.", "University pages must show prerequisites, GPA, tests, materials, deadlines, costs, and curriculum clearly.", "Better university requirement briefs and source-confidence labels."),
        ("18. Product readiness", "Many things are not working. Make it final product ready.", "Stability, packaging, and demo confidence became the priority.", "Build/release checks, cleaner routes, README, local run instructions."),
        ("19. Code and routes", "I need all code file. Can you give its route?", "The product needed to be transferable, runnable, and explainable.", "Code zip, route summary, startup instructions."),
        ("20. B2B polish", "Apply the Hangzhou Hustle Product Polish Blueprint.", "Package the app as premium B2B software, not a student prototype.", "Enterprise badge, white-label engine, print report, locked premium modules, live diagnostic sliders."),
        ("21. Origin document", "Make a list from how we started this in a document: all the prompts and the motive behind it.", "Capture the project story so the product has a clear strategic memory.", "This document."),
    ]
    add_timeline_table(doc, timeline)

    add_heading(doc, "3. Motive Map", 1)
    motive_rows = [
        ("Personal AI Twin", "Aashi Dreams should understand Anushka better over time and turn her profile into decisions."),
        ("Admissions Intelligence", "Every university needs rank, acceptance difficulty, requirements, roadmap, fit score, and missing-profile delta."),
        ("Scholarship Survival", "Scholarship tracking must be deep because funding is competitive and time-sensitive."),
        ("Research Reputation", "The gaming/derealization paper and future research ideas are major differentiators."),
        ("Career and ROI", "The app should compare degrees, countries, visa routes, salaries, work authorization, and career outcomes."),
        ("Network Intelligence", "Professors, alumni, admissions officers, researchers, and product managers become trackable assets."),
        ("B2B Packaging", "The same engine can become a white-label product for coaching centers if it looks trustworthy and official."),
    ]
    add_key_value_table(doc, motive_rows)

    add_heading(doc, "4. Final Product Positioning", 1)
    add_body(
        doc,
        "Aashi Dreams is best positioned as an AI operating system for a student's future. It should not be described as only a study-abroad platform, application tracker, or career guidance app. Its stronger promise is: predict, plan, optimize, and execute the path from student profile to university, scholarship, career, and global mobility.",
    )
    for item in [
        "For Anushka: a private strategic command center for admissions, scholarships, research, and career planning.",
        "For coaching centers: a white-label diagnostic engine that makes counseling look data-rich, premium, and parent-ready.",
        "For future expansion: a living Student Digital Twin that gets smarter as more evidence, results, tasks, and opportunities are added.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5. One-Line Origin Story", 1)
    add_body(
        doc,
        "Aashi Dreams started with the refusal to build another ordinary admissions tracker and became a personal AI operating system for turning Anushka's psychology, research, startup, scholarships, university choices, and career ambitions into one evolving strategy engine.",
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Aashi Dreams Project Record")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
